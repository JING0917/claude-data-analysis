#!/usr/bin/env python3
import docx
import sys

def read_docx(file_path):
    """读取Word文档内容"""
    doc = docx.Document(file_path)

    print(f"文档段落数: {len(doc.paragraphs)}")
    print(f"文档表格数: {len(doc.tables)}")
    print("=" * 70)

    # 读取所有段落
    print("文档内容:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # 只显示非空段落
            print(f"[段落 {i}] {para.text[:100]}...")

    print("\n" + "=" * 70)
    print("表格内容:")

    # 读取所有表格
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx}:")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            print(f"  行 {row_idx}: {row_data}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python read_docx.py <docx文件路径>")
        sys.exit(1)

    file_path = sys.argv[1]
    read_docx(file_path)