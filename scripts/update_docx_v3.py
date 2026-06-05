#!/usr/bin/env python3
import docx
import sys

def update_docx(input_path, output_path):
    """更新Word文档的结论和数据"""
    doc = docx.Document(input_path)

    print(f"正在更新文档: {input_path}")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")

    # 1. 更新结论段落 (段落2-13)
    new_conclusion_paragraphs = [
        "在美团仅采集第一页订单的技术限制下，分析590,761位采集美团下单小蚕用户发现：",
        "",
        "【核心发现】通过结合匹配类型的用户分层分析，精准识别41,468名最高价值用户（占7.02%）：",
        "1. 最高价值用户：41,468名“用户自己下单的高潜力用户”，平均美团订单6.33单，美团比小蚕多4.42单（差异率69.98%）",
        "2. 数据可靠性：405,447用户（68.63%）数据逻辑一致（小蚕美团订单≤采集美团订单），分析结果可信",
        "3. 分层改进：避免了过去4,255名“订单完全由其他有关系用户下单”用户误判为高潜力用户的问题",
        "",
        "【关键洞察】",
        "• 用户自己下单的高潜力用户（41,468名）是转化价值最高的目标群体，推广ROI预期最优",
        "• 用户美团数据可能异常用户：105,926名（17.93%），数据质量问题需要关注",
        "• 负差异用户：185,314名（31.37%），受数据采集限制（仅采集第一页订单）影响",
        "",
        "【运营优先级】",
        "• 最高优先级：41,468名用户自己下单的高潜力用户",
        "• 高优先级：45,448名用户（用户自己下单的中潜力用户+其他匹配类型的高潜力用户）",
        "• 中优先级：70,317名用户",
        "• 低优先级：151,508名无差异用户",
        "• 监控即可：282,020名负差异或数据异常用户",
    ]

    # 找到结论部分的起始和结束位置
    # 根据读取的内容，段落2是结论开始，段落13是结论结束前
    # 但为了安全，我们找到包含"结论"的段落
    start_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if "结论" in para.text and i < len(doc.paragraphs) - 1:
            start_idx = i + 1  # 从下一个段落开始
            break

    if start_idx != -1:
        # 找到结论结束位置（下一个标题或空段落）
        end_idx = start_idx
        for i in range(start_idx, len(doc.paragraphs)):
            if i >= len(doc.paragraphs):
                break
            # 如果遇到下一个标题或特定标记，停止
            if "口径" in doc.paragraphs[i].text or i == start_idx + 20:  # 最多20段
                end_idx = i
                break

        print(f"更新结论段落: 从{start_idx}到{end_idx}")

        # 删除旧的结论段落
        for i in range(end_idx - 1, start_idx - 1, -1):
            if i < len(doc.paragraphs):
                # 清空段落内容而不是删除，保持格式
                p = doc.paragraphs[i]
                p.clear()
                # 添加新内容（如果需要）

        # 在起始位置插入新结论
        for i, text in enumerate(new_conclusion_paragraphs):
            if start_idx + i < len(doc.paragraphs):
                doc.paragraphs[start_idx + i].text = text
            else:
                # 如果段落不够，添加新段落
                doc.add_paragraph(text)

    # 2. 更新表格数据
    # 表格3: 用户分层结果 - 需要更新为结合匹配类型的数据？
    # 表格4: 匹配类型分布 - 数据应该还是准确的

    # 检查表格3（用户分层结果）
    if len(doc.tables) >= 4:  # 表格索引从0开始
        table3 = doc.tables[3]  # 第四个表格
        print(f"表格3行数: {len(table3.rows)}")

        # 这是用户分层结果表格，数据可能需要更新
        # 但用户说微调，所以暂时保持原数据或小范围更新
        # 这里可以更新部分数据

    # 保存更新后的文档
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")

    # 验证更新
    print("\n更新后验证:")
    updated_doc = docx.Document(output_path)
    for i, para in enumerate(updated_doc.paragraphs[:15]):  # 只显示前15段
        if para.text.strip():
            print(f"[{i}] {para.text[:80]}...")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python update_docx_v3.py <输入文件> <输出文件>")
        print("示例: python update_docx_v3.py analysis_reports/小蚕用户美团下单频次分析V3.docx analysis_reports/小蚕用户美团下单频次分析V3_更新版.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    update_docx(input_file, output_file)