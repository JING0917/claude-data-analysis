#!/usr/bin/env python3
"""
简洁地更新Word文档，只修改必要部分
"""
import docx
import sys

def clean_update_docx(input_path, output_path):
    doc = docx.Document(input_path)

    print(f"处理文档: {input_path}")

    # 1. 找到"结论"段落和"口径"段落
    conclusion_start = -1
    conclusion_end = -1

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "结论" in text and conclusion_start == -1:
            conclusion_start = i
        elif "口径" in text and conclusion_start != -1 and conclusion_end == -1:
            conclusion_end = i
            break

    if conclusion_start == -1:
        print("未找到'结论'部分")
        return

    if conclusion_end == -1:
        # 如果没有找到"口径"，则用段落数
        conclusion_end = len(doc.paragraphs)

    print(f"结论部分: 段落{conclusion_start}到{conclusion_end}")

    # 2. 清空结论部分（从conclusion_start+1到conclusion_end-1）
    for i in range(conclusion_end - 1, conclusion_start, -1):
        if i < len(doc.paragraphs):
            doc.paragraphs[i].clear()

    # 3. 在结论标题后插入新内容
    new_content = [
        "在美团仅采集第一页订单的技术限制下，分析590,761位采集美团下单小蚕用户发现：",
        "",
        "【核心发现】通过结合匹配类型的用户分层分析，精准识别41,468名最高价值用户（占7.02%）：",
        "1. 最高价值用户：41,468名“用户自己下单的高潜力用户”，平均美团订单6.33单，美团比小蚕多4.42单（差异率69.98%）",
        "2. 数据可靠性：405,447用户（68.63%）数据逻辑一致（小蚕美团订单≤采集美团订单），分析结果可信",
        "3. 分层改进：避免了过去4,255名“订单完全由其他有关系用户下单”用户误判为高潜力用户的问题",
        "",
        "【关键洞察】",
        "• 用户自己下单的高潜力用户（41,468名）是转化价值最高的目标群体，推广ROI预期最优",
        "• 用户美团数据可能异常用户：105,926名（17.93%），数据质量问题需要关注",
        "• 负差异用户：185,314名（31.37%），受数据采集限制（仅采集第一页订单）影响",
        "",
        "【运营优先级】",
        "• 最高优先级：41,468名用户自己下单的高潜力用户",
        "• 高优先级：45,448名用户（用户自己下单的中潜力用户+其他匹配类型的高潜力用户）",
        "• 中优先级：70,317名用户",
        "• 低优先级：151,508名无差异用户",
        "• 监控即可：282,020名负差异或数据异常用户",
    ]

    # 插入新内容
    insert_pos = conclusion_start + 1
    for i, text in enumerate(new_content):
        if insert_pos + i < len(doc.paragraphs):
            doc.paragraphs[insert_pos + i].text = text
        else:
            # 添加新段落
            doc.add_paragraph(text)

    # 4. 更新运营建议部分（包含"83,148"的段落）
    for i, para in enumerate(doc.paragraphs):
        if "83,148" in para.text:
            print(f"更新运营建议段落 {i}: {para.text[:50]}...")
            # 替换为新的运营建议
            new_text = "存量用户提升：针对41,468名用户自己下单的高潜力用户开展精准推广，这些用户转化价值最高，ROI预期最优。"
            para.text = new_text
            break

    # 5. 保存文档
    doc.save(output_path)
    print(f"文档已保存: {output_path}")

    # 验证
    print("\n验证更新:")
    updated = docx.Document(output_path)
    for i, para in enumerate(updated.paragraphs):
        if i <= conclusion_start + 20 or "41,468" in para.text:
            if para.text.strip():
                print(f"[{i}] {para.text[:80]}...")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python clean_update_docx.py <输入文件> <输出文件>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    clean_update_docx(input_file, output_file)